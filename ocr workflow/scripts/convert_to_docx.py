import os
import sys
import re
import argparse
import docx
from docx import Document
from docx.shared import Pt, Inches

def parse_inline_formatting(paragraph, text):
    """
    Parses basic bold (**) and italic (*) formatting from Markdown and appends runs.
    """
    # Split text by bold markers (**) and italic markers (*)
    # Pattern captures bold '**text**' or italic '*text*' or normal text.
    pattern = re.compile(r'(\*\*.*?\*\*|\*.*?\*|[^\*\s]+(?:\s+[^\*\s]+)*|\s+)')
    tokens = pattern.findall(text)
    
    for token in tokens:
        if token.startswith('**') and token.endswith('**'):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith('*') and token.endswith('*'):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)

def markdown_to_docx(md_path: str, docx_path: str) -> bool:
    """
    Reads a Markdown file and converts it into a Word Document (.docx).
    """
    if not os.path.exists(md_path):
        print(f"Error: Markdown file not found at {md_path}")
        return False
        
    doc = Document()
    
    # Configure default style margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
        
    in_code_block = False
    
    for line in lines:
        line_stripped = line.strip()
        
        # Handle code block start/end
        if line_stripped.startswith("```"):
            in_code_block = not in_code_block
            continue
            
        if in_code_block:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(line)
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            continue
            
        # Handle Headings
        match_heading = re.match(r'^(#{1,6})\s+(.*)$', line_stripped)
        if match_heading:
            level = len(match_heading.group(1))
            text = match_heading.group(2)
            # Remove inline md formatting from heading
            text_clean = re.sub(r'\*\*|__|\*|_', '', text)
            doc.add_heading(text_clean, level=level)
            continue
            
        # Handle Page Separators
        if line_stripped.startswith("--- Page ") or line_stripped.startswith("--- Page"):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run(line_stripped)
            run.bold = True
            run.font.size = Pt(11)
            continue
            
        # Handle horizontal separators
        if line_stripped == "---" or line_stripped == "***":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            run = p.add_run("__________________________________________________")
            run.font.color.rgb = docx.shared.RGBColor(128, 128, 128)
            continue
            
        # Handle bullet lists
        match_bullet = re.match(r'^[\-\*\+]\s+(.*)$', line_stripped)
        if match_bullet:
            text = match_bullet.group(1)
            p = doc.add_paragraph(style='List Bullet')
            parse_inline_formatting(p, text)
            continue
            
        # Handle numbered lists
        match_number = re.match(r'^(\d+)\.\s+(.*)$', line_stripped)
        if match_number:
            text = match_number.group(2)
            p = doc.add_paragraph(style='List Number')
            parse_inline_formatting(p, text)
            continue
            
        # Handle empty line
        if not line_stripped:
            continue
            
        # Regular paragraph
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        parse_inline_formatting(p, line_stripped)
        
    doc.save(docx_path)
    print(f"Success: Converted '{md_path}' to '{docx_path}'")
    return True

if __name__ == "__main__":
    # Force UTF-8 for stdout to prevent Windows encoding errors
    if sys.stdout.encoding != 'utf-8':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    parser = argparse.ArgumentParser(description="Convert Markdown to Word Document (.docx).")
    parser.add_argument("input_md", help="Path to the input Markdown file.")
    parser.add_argument("output_docx", help="Path to the output Word Document (.docx).")
    
    args = parser.parse_args()
    
    success = markdown_to_docx(args.input_md, args.output_docx)
    sys.exit(0 if success else 1)
